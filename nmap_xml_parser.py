import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Cm, Pt
import sys
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.shared import qn  
from docx.oxml.xmlchemy import OxmlElement
import argparse

# HANDLE ARGUMENTS -- START
parser = argparse.ArgumentParser()
parser.add_argument("-f", "--file", type=str, help="file to be parsed, ex: nmap_scan.xml")
parser.add_argument("-o", "--output", type=str, help="output file, ex: ports_services")

args = parser.parse_args()
# HANDLE ARGUMENTS -- END

if args.output is not None:
    output_file = args.output
else:
    print("No output file specified!")
    sys.exit(1)

if not os.path.exists(args.file):
    print("The specified file does not exist! Please ensure the path and the file name are correct.")
    sys.exit(1)
else:
    file_to_read = args.file
    with open(file_to_read, "r") as file:
        html_doc = file.read()

soup = BeautifulSoup(html_doc, "html.parser")

class Parse():
    def __init__(self, soup):
        self.soup = soup
        self.host_ips = []
        self.ports_services = []
    
    def get_hosts(self):
        hosts_tags = self.soup.find_all("host")
        
        for host in hosts_tags:
            self.parse_tags(host)
        
        self.display_results()

    def parse_tags(self, host):
        host_ip = host.find("address").get("addr")
        port_tags = host.find_all("port")
        # service_tags = host.find_all("service")
        
        ports_services = []
        for port in port_tags:
            port_id = port.get("portid")
            
            if port.find("service"):
                service_tag = port.find("service")
                service_name = service_tag.get("name")
            else:
                service_name = "unknown"
            
            ports_services.append([port_id, service_name])

        self.append_values(host_ip, ports_services)
    
    def append_values(self, ip, ports_services):
        self.host_ips.append(ip)
        self.ports_services.append(ports_services)
    
    def display_results(self):
        word_document = Document()
        document_name = output_file
        
        table = word_document.add_table(0, 0) # we add rows iteratively
        table.style = 'TableGrid'
        table.style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        table.style.paragraph_format.line_spacing = Pt(14)
        
        first_column_width = 20
        second_column_width = 6
        third_column_width = 15
        table.add_column(Cm(first_column_width))
        table.add_column(Cm(second_column_width))
        table.add_column(Cm(third_column_width))
        
        headers = ["Host", "Ports", "Services"]
        table.add_row()
        
        self._set_cell_background(table.rows[0].cells[0], 'D9D9D9')
        self._set_cell_background(table.rows[0].cells[1], 'D9D9D9')
        self._set_cell_background(table.rows[0].cells[2], 'D9D9D9')
        table.rows[0].height = Cm(0.9)
        
        for i in range(len(headers)):
            cell= table.cell(0, i)
            cell.text = headers[i]
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.name = "Arial"
            cell.paragraphs[0].runs[0].font.size = Pt(12)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        index = 1
        number_of_hosts = len(self.host_ips)
        while (index < number_of_hosts + 1):
            table.add_row()
            row = table.rows[index]
            row.height = Cm(0.7)
            
            ip = self.host_ips[index - 1]
            ports_services = self.ports_services[index - 1]
            print("-" * 100)
            print("IP ADDRESS: " + str(ip))
            print("PORT" + " " * 5 + "SERVICE")
            for element in ports_services:
                print(element[0] + " " * 5 + element[1])
            
            ports = []
            services = []
            for element in ports_services:
                ports.append(element[0])
                services.append(element[1])
            
            ports_str = '\n'.join([str(elem) for elem in ports])
            services_str = '\n'.join([str(elem) for elem in services])
            
            attrs = [ip, ports_str, services_str]
            
            for i in range(len(attrs)):
                cell = row.cells[i]
                cell.text = str(attrs[i])
                cell.paragraphs[0].runs[0].font.name = "Arial"
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # row.cells[0].text = str(ip)
            # row.cells[1].text = str(ports_str)
            # row.cells[2].text = str(services_str)
            index += 1
        word_document.add_page_break()
        word_document.save(document_name + '.docx')
        
    def _set_cell_background(self, cell, fill, colour=None, val=None):
        """
        @fill: Specifies the colour to be used for the background
        @colour: Specifies the colour to be used for any foreground
        pattern specified with the val attribute
        @val: Specifies the pattern to be used to lay the pattern
        colour over the background colour.
        """
        
        cell_properties = cell._element.tcPr
        try:
            cell_shading = cell_properties.xpath('w:shd')[0]  # in case there's already shading
        except IndexError:
            cell_shading = OxmlElement('w:shd') # add new w:shd element to it
        
        if fill:
            cell_shading.set(qn('w:fill'), fill)  # set fill property, respecting namespace
        if colour:
            pass # TODO
        if val:
            pass # TODO
        cell_properties.append(cell_shading)  # finally extend cell props with shading element


my_project = Parse(soup)
my_project.get_hosts()